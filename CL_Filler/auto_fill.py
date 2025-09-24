import os
from docx import Document
import re
from datetime import datetime

# ========== CONFIGURAÇÕES DE CAMINHOS ==========
# Pasta do arquivo modelo
PASTA_MODELO = r"C:\Users\SeuUsuario\Documents\Modelos"
# Nome do arquivo modelo
NOME_ARQUIVO_MODELO = "CL RESID LOCADOR X LOCATÁRIO - MODELO.docx"
# Pasta onde vai ser salvo os contratos
PASTA_DESTINO = r"C:\Users\SeuUsuario\Documents\Contratos_Gerados"


# Caminhos completos (não alterar)
CAMINHO_MODELO = os.path.join(PASTA_MODELO, NOME_ARQUIVO_MODELO)

# Criar pasta de destino se não existir
if not os.path.exists(PASTA_DESTINO):
    os.makedirs(PASTA_DESTINO)
    print(f"Pasta de destino criada: {PASTA_DESTINO}")

def extrair_nome_sobrenome(nome_completo):
    """Extrai apenas o primeiro nome e último sobrenome"""
    partes = nome_completo.strip().split()
    if len(partes) >= 2:
        return f"{partes[0]} {partes[-1]}"
    return partes[0] if partes else ""

def coletar_dados_usuario():
    """Coleta todos os dados necessários do usuário"""
    print("=== GERADOR DE CONTRATO DE LOCAÇÃO ===\n")
    
    dados = {}
    
    # Dados do Locador
    print("--- DADOS DO LOCADOR ---")
    dados['LOCADOR'] = input("Nome completo do locador: ").strip()
    dados['NACIO1'] = input("Nacionalidade: ").strip()
    dados['ESTADOC1'] = input("Estado civil: ").strip()
    dados['PROF1'] = input("Profissão: ").strip()
    dados['CPF1'] = input("CPF (apenas números): ").strip()
    dados['LOC1'] = input("Endereço completo: ").strip()
    
    print("\n--- DADOS DO LOCATÁRIO ---")
    dados['LOCATÁRIO'] = input("Nome completo do locatário: ").strip()
    dados['NACIO2'] = input("Nacionalidade: ").strip()
    dados['ESTADOC2'] = input("Estado civil: ").strip()
    dados['PROF2'] = input("Profissão: ").strip()
    dados['CPF2'] = input("CPF (apenas números): ").strip()
    dados['LOC2'] = input("Endereço completo: ").strip()
    
    # Verificar se haverá fiador
    tem_fiador = input("\nTerá fiador? (s/n): ").strip().lower() == 's'
    
    if tem_fiador:
        print("\n--- DADOS DO FIADOR ---")
        dados['FIADOR'] = input("Nome completo do fiador: ").strip()
        dados['NACIO3'] = input("Nacionalidade: ").strip()
        dados['ESTADOC3'] = input("Estado civil: ").strip()
        dados['PROF3'] = input("Profissão: ").strip()
        dados['CPF3'] = input("CPF (apenas números): ").strip()
        dados['LOC3'] = input("Endereço completo: ").strip()
    else:
        # Se não há fiador, remover as seções relacionadas
        dados['FIADOR'] = ""
        dados['NACIO3'] = ""
        dados['ESTADOC3'] = ""
        dados['PROF3'] = ""
        dados['CPF3'] = ""
        dados['LOC3'] = ""
    
    # Dados do Imóvel
    print("\n--- DADOS DO IMÓVEL ---")
    dados['IMOVEL'] = input("Descrição completa do imóvel (endereço, tipo, etc.): ").strip()
    
    # Dados do Contrato
    print("\n--- DADOS DO CONTRATO ---")
    dados['PRAZO'] = input("Prazo em meses: ").strip()
    
    # Data de início
    while True:
        try:
            data_inicio = input("Data de início (DD/MM/AAAA): ").strip()
            datetime.strptime(data_inicio, "%d/%m/%Y")
            dados['DATACL'] = data_inicio
            dados['DATA_CL'] = data_inicio  # Para a data de assinatura
            break
        except ValueError:
            print("Formato de data inválido. Use DD/MM/AAAA")
    
    # Valor do aluguel
    while True:
        try:
            valor = float(input("Valor do aluguel (apenas números, ex: 1200.50): ").strip().replace(',', '.'))
            dados['VALOR'] = f"{valor:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            # Converter valor para extenso (simplificado)
            dados['VALOR_DIG'] = input("Valor por extenso (ex: mil e duzentos reais): ").strip()
            break
        except ValueError:
            print("Valor inválido. Digite apenas números.")
    
    # Cláusula 18 (Garantias)
    print("\n--- GARANTIAS LOCATÍCIAS ---")
    print("Cole o texto completo da cláusula de garantia:")
    print("(Pressione Enter duas vezes consecutivas para finalizar)")
    print("-" * 50)
    
    linhas_clausula = []
    linhas_vazias_consecutivas = 0
    
    while True:
        linha = input()
        
        if linha.strip() == "":
            linhas_vazias_consecutivas += 1
            if linhas_vazias_consecutivas >= 2:
                break
        else:
            linhas_vazias_consecutivas = 0
            linhas_clausula.append(linha)
    
    dados['CLAUSULA18'] = "\n".join(linhas_clausula)
    
    return dados, tem_fiador

def processar_documento(arquivo_origem, dados, tem_fiador):
    """Processa o documento substituindo as variáveis pelos dados fornecidos"""
    try:
        # Carregar o documento
        doc = Document(arquivo_origem)
        
        # Processar parágrafos
        for paragraph in doc.paragraphs:
            texto_original = paragraph.text
            texto_modificado = texto_original
            
            # Substituir todas as variáveis
            for chave, valor in dados.items():
                placeholder = f"[{chave}]"
                if placeholder in texto_modificado:
                    texto_modificado = texto_modificado.replace(placeholder, str(valor))
            
            # Se não há fiador, remover seções relacionadas
            if not tem_fiador:
                # Remover linha do fiador no preâmbulo se estiver entre colchetes
                if "[ FIADOR (A):" in texto_modificado:
                    # Encontrar e remover toda a seção do fiador
                    pattern = r'\*\*\[\s*FIADOR \(A\):.*?\]\*\*'
                    texto_modificado = re.sub(pattern, '', texto_modificado, flags=re.DOTALL)
            
            # Atualizar o parágrafo se houve mudanças
            if texto_modificado != texto_original:
                paragraph.clear()
                paragraph.add_run(texto_modificado)
        
        # Processar tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_original = cell.text
                    texto_modificado = texto_original
                    
                    for chave, valor in dados.items():
                        placeholder = f"[{chave}]"
                        if placeholder in texto_modificado:
                            texto_modificado = texto_modificado.replace(placeholder, str(valor))
                    
                    if texto_modificado != texto_original:
                        cell.text = texto_modificado
        
        # Processar seções de assinatura
        if not tem_fiador:
            # Esta parte pode precisar de ajustes dependendo da estrutura do documento
            for paragraph in doc.paragraphs:
                if "FIADOR:" in paragraph.text:
                    # Remover ou modificar as linhas de assinatura do fiador
                    if "ANOR CARDOZO PEREIRA" in paragraph.text or "NATALINA DE MOURA PEREIRA" in paragraph.text:
                        paragraph.clear()
        
        return doc
        
    except Exception as e:
        print(f"Erro ao processar documento: {str(e)}")
        return None

def main():
    # Exibir configurações atuais
    print("=== CONFIGURAÇÕES DE CAMINHOS ===")
    print(f"Pasta do modelo: {PASTA_MODELO}")
    print(f"Arquivo modelo: {NOME_ARQUIVO_MODELO}")
    print(f"Pasta de destino: {PASTA_DESTINO}")
    print("-" * 50)
    
    # Verificar se o arquivo modelo existe
    if not os.path.exists(CAMINHO_MODELO):
        print(f"ERRO: Arquivo modelo não encontrado!")
        print(f"Caminho esperado: {CAMINHO_MODELO}")
        print("\nPara corrigir:")
        print("1. Altere as variáveis PASTA_MODELO e NOME_ARQUIVO_MODELO no início do código")
        print("2. Ou mova o arquivo modelo para o caminho indicado acima")
        return
    
    # iniciar coleta
    dados, tem_fiador = coletar_dados_usuario()
    
    print("\nProcessando contrato...")
    doc_processado = processar_documento(CAMINHO_MODELO, dados, tem_fiador)
    
    if doc_processado is None:
        print("Erro ao processar o documento!")
        return
    
    # Gerar nome do arquivo
    nome_locador = extrair_nome_sobrenome(dados['LOCADOR'])
    nome_locatario = extrair_nome_sobrenome(dados['LOCATÁRIO'])
    nome_arquivo = f"CL RESID {nome_locador} X {nome_locatario}.docx"
    
    # Caminho pra salvar o arquivo
    caminho_arquivo_saida = os.path.join(PASTA_DESTINO, nome_arquivo)
    try:
        doc_processado.save(caminho_arquivo_saida)
        print(f"\nContrato gerado com sucesso!")
        print(f"Arquivo salvo como: {nome_arquivo}")
        print(f"Local: {caminho_arquivo_saida}")
        
    except Exception as e:
        print(f"Erro ao salvar arquivo: {str(e)}")

if __name__ == "__main__":
    main()
