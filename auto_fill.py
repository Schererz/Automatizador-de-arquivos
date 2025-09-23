import os
from pathlib import Path

# Importações para DOCX e PDF
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx2pdf import convert
except ImportError:
    print("ERRO: Instale as dependências primeiro!")
    print("Execute: pip install python-docx docx2pdf")
    exit()

def coletar_dados():
    """Coleta as informações necessárias do usuário."""
    print("Por favor, insira os dados do contrato:")
    dados = {}
    
    # Dados do Locador
    print("\n--- Dados do Locador ---")
    dados['nome_locador'] = input("Nome do Locador: ").upper()
    dados['nacionalidade1'] = input("Nacionalidade do Locador: ")
    dados['estadocivil1'] = input("Estado Civil do Locador: ")
    dados['profissao1'] = input("Profissão do Locador: ")
    dados['cpf1'] = input("CPF do Locador: ")
    dados['localizacao1'] = input("Localização do Locador (Rua, nº, Bairro, Cidade/Estado, CEP): ").upper()

    # Dados do Locatário
    print("\n--- Dados do Locatário ---")
    dados['nome_locatario'] = input("Nome do Locatário: ").upper()
    dados['nacionalidade2'] = input("Nacionalidade do Locatário: ")
    dados['estadocivil2'] = input("Estado Civil do Locatário: ")
    dados['profissao2'] = input("Profissão do Locatário: ")
    dados['cpf2'] = input("CPF do Locatário: ")
    dados['localizacao2'] = input("Localização do Locatário (Rua, nº, Bairro, Cidade/Estado, CEP): ").upper()

    # Dados do Imóvel
    print("\n--- Dados do Imóvel ---")
    dados['localizacao_imovel'] = input("Localização do Imóvel (Rua, nº, Bairro, Cidade/Estado, CEP): ").upper()

    # Outros dados
    print("\n--- Outros Dados ---")
    dados['data_contrato'] = input("Data do Contrato (ex: 20 de outubro de 2025): ")
    
    return dados

def preencher_e_salvar_docx(modelo_path, dados):
    """
    Preenche o modelo DOCX com os dados, e salva a cópia preenchida.
    Retorna o caminho do novo arquivo DOCX.
    """
    try:
        documento = docx.Document(modelo_path)
    except Exception as e:
        print(f"Erro ao abrir o modelo DOCX: {e}")
        return None

    # Mapeamento dos dados para os marcadores no documento
    # Marcadores sem espaçamento para corresponder ao novo modelo
    marcadores = {
        '[NOMELOCADOR]': dados['nome_locador'],
        '[NACIONALIDADE1]': dados['nacionalidade1'],
        '[ESTADOCIVIL1]': dados['estadocivil1'],
        '[PROFISSÃO1]': dados['profissao1'],
        '[CPF1]': dados['cpf1'],
        '[LOCALIZAÇÃO1]': dados['localizacao1'],
        '[NOMELOCATÁRIO]': dados['nome_locatario'],
        '[NACIONALIDADE2]': dados['nacionalidade2'],
        '[ESTADOCIVIL2]': dados['estadocivil2'],
        '[PROFISSÃO2]': dados['profissao2'],
        '[CPF2]': dados['cpf2'],
        '[LOCALIZAÇÃO2]': dados['localizacao2'],
        '[LOCALIZAÇÃOIMÓVEL]': dados['localizacao_imovel'],
        '[DATA DO CONTRATO]': dados['data_contrato']
    }

    # Lógica de substituição melhorada para ser mais robusta
    def substituir_no_texto(texto):
        for key, value in marcadores.items():
            texto = texto.replace(key, value)
        return texto

    for p in documento.paragraphs:
        p.text = substituir_no_texto(p.text)
        
    for table in documento.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = substituir_no_texto(p.text)
    
    # Cria o novo caminho para o arquivo preenchido
    pasta_saida = Path(modelo_path).parent
    novo_nome_base = f"Contrato - {dados['nome_locatario']} - {dados['data_contrato']}"
    caminho_saida_docx = pasta_saida / f"{novo_nome_base}.docx"
    
    # Evita sobrescrever arquivos existentes
    contador = 1
    while caminho_saida_docx.exists():
        caminho_saida_docx = pasta_saida / f"{novo_nome_base} ({contador}).docx"
        contador += 1

    try:
        documento.save(caminho_saida_docx)
        print(f"✅ Documento DOCX salvo com sucesso em: {caminho_saida_docx.name}")
        return caminho_saida_docx
    except Exception as e:
        print(f"❌ Erro ao salvar o documento: {e}")
        return None

def converter_para_pdf(caminho_docx):
    """Converte o arquivo DOCX preenchido para PDF."""
    if not caminho_docx or not Path(caminho_docx).exists():
        print("❌ Caminho do arquivo DOCX inválido para conversão.")
        return

    try:
        caminho_saida_pdf = caminho_docx.with_suffix('.pdf')
        convert(str(caminho_docx), str(caminho_saida_pdf))
        print(f"✅ Conversão para PDF concluída com sucesso: {caminho_saida_pdf.name}")
    except Exception as e:
        print(f"❌ Erro ao converter para PDF: {e}")

def main():
    """Função principal que orquestra o programa."""
    print("🔄 PREENCHEDOR AUTOMÁTICO DE CONTRATOS")
    print("=" * 40)
    
    # CONFIRA O CAMINHO DO SEU ARQUIVO DE MODELO AQUI
    modelo_path = "modelo.docx"
    
    # 1. Coleta os dados do usuário
    dados_contrato = coletar_dados()
    
    # 2. Preenche e salva o novo arquivo DOCX
    caminho_docx_pronto = preencher_e_salvar_docx(modelo_path, dados_contrato)
    
    # 3. Converte para PDF
    if caminho_docx_pronto:
        converter_para_pdf(caminho_docx_pronto)
        
    print("\n✅ Processamento concluído!")
    input("Pressione Enter para sair...")

if __name__ == "__main__":
    main()